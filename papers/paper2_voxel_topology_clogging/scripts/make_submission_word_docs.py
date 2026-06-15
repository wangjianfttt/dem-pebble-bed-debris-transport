#!/usr/bin/env python3
"""Create Word-format submission support documents for Paper 2."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[3]
PAPER_DIR = PROJECT_ROOT / "papers" / "paper2_voxel_topology_clogging"
SUBMISSION_DIR = PAPER_DIR / "submission"
FIGURE_DIR = PAPER_DIR / "figures"
OUT_DIR = SUBMISSION_DIR / "word"


TITLE = "Separating breakthrough from pore-network clogging indicators in gas-driven fine-debris transport through Li4SiO4 pebble beds"
DOI = "10.5281/zenodo.20699272"
REPOSITORY = "https://github.com/wangjianfttt/dem-pebble-bed-debris-transport"


def style_document(doc: Document) -> None:
    """Apply a restrained Word style suitable for Elsevier submission documents."""
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, size, color in [
        ("Heading 1", 16, RGBColor(0x1F, 0x4D, 0x78)),
        ("Heading 2", 13, RGBColor(0x1F, 0x4D, 0x78)),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True


def add_title(doc: Document, title: str, subtitle: str | None = None) -> None:
    """Add a clean title block without relying on Word's built-in Title style."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(title)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(15)
    run.font.color.rgb = RGBColor(0x11, 0x11, 0x11)
    if subtitle:
        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = sub.add_run(subtitle)
        r.font.name = "Arial"
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def add_bullet(doc: Document, text: str) -> None:
    """Add one Word bullet item."""
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.space_after = Pt(3)
    para.add_run(text)


def build_cover_letter(out_path: Path) -> Path:
    """Create the cover letter in Word format."""
    doc = Document()
    style_document(doc)
    add_title(doc, "Cover Letter", "Chemical Engineering Science submission")

    paragraphs = [
        "Dear Editor,",
        f'We are pleased to submit our manuscript entitled "{TITLE}" for consideration in Chemical Engineering Science.',
        "This manuscript addresses a diagnostic problem in gas-purged packed beds: fine debris may reach the outlet before the pore network shows measurable structural or hydraulic degradation. We examine this issue in a Li4SiO4 ceramic breeder pebble bed, where debris migration, retention and purge-pathway degradation must be interpreted carefully.",
        "The work combines DEM calculations with DEM-derived pore reconstruction and bounded pressure checks to distinguish four related but non-equivalent responses: outlet breakthrough, local deposition, connected-pore degradation and hydraulic confirmation. The central contribution is a bounded workflow for separating transport arrival from pore-network clogging indicators, rather than a universal clogging law or pressure-calibrated safety criterion.",
        "The manuscript is organized around three mechanism axes in the sampled cases: drag-to-weight ratio primarily affects downstream penetration, localized internal release produces retained-bulk/sparse-front separation, and debris inventory controls local deposition intensity. The response-landscape analysis explains why these observables should not be collapsed into a single clogging coordinate.",
        "We believe the manuscript fits Chemical Engineering Science because it focuses on transport and structure-response mechanisms in a packed granular medium. The manuscript explicitly avoids claiming experimental validation, universal transition laws or pressure-calibrated safety limits. It instead provides a reproducible numerical framework for interpreting fine-debris transport and clogging indicators within the studied parameter space and present DEM assumptions.",
        f"Processed data, figure source tables and scripts are available in the public GitHub-Zenodo repository {REPOSITORY}, archived at DOI {DOI}. Raw DEM dump and restart files are not included in the lightweight public deposit because of their size and are available from the corresponding author upon reasonable request.",
        "We confirm that this manuscript is original, has not been published elsewhere and is not under consideration by another journal.",
        "Sincerely,",
        "Jian Wang\non behalf of all authors",
    ]
    for text in paragraphs:
        doc.add_paragraph(text)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    return out_path


def build_highlights(out_path: Path) -> Path:
    """Create highlights in Word format."""
    doc = Document()
    style_document(doc)
    add_title(doc, "Highlights", "Chemical Engineering Science submission")
    highlights = [
        "DEM resolves fines transport in a stiff gas-purged Li4SiO4 pebble bed.",
        "Topology metrics separate breakthrough from pore-network degradation in the sampled cases.",
        "Drive, source release and inventory activate different debris observables.",
        "Discrete coverage mapping bounds evidence without phase-map claims.",
        "Cropped-domain pressure checks support a bounded pre-clogging interpretation.",
    ]
    for item in highlights:
        add_bullet(doc, item)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    return out_path


def build_graphical_abstract_doc(out_path: Path) -> Path:
    """Create a Word document containing the graphical abstract and caption."""
    doc = Document()
    style_document(doc)
    add_title(doc, "Graphical Abstract", "Chemical Engineering Science submission")
    figure_path = FIGURE_DIR / "paper2_graphical_abstract.png"
    doc.add_picture(str(figure_path), width=Inches(6.5))
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption.add_run(
        "Graphical abstract. DEM trajectories, pore reconstruction and bounded pressure checks provide complementary indicators for interpreting fines breakthrough, local deposition, connected-pore loss and hydraulic confirmation in the sampled Li4SiO4 pebble-bed cases."
    )
    run.italic = True
    run.font.size = Pt(9)

    note = doc.add_paragraph()
    note.add_run("Scope note: ").bold = True
    note.add_run(
        "The workflow is bounded to the studied parameter space, finite voxel resolution and present one-way DEM assumptions; it is not a universal transition law or pressure-calibrated safety criterion."
    )
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    return out_path


def main() -> int:
    """Generate all Word-format submission support documents."""
    outputs = [
        build_cover_letter(OUT_DIR / "cover_letter_CES.docx"),
        build_highlights(OUT_DIR / "highlights_CES.docx"),
        build_graphical_abstract_doc(OUT_DIR / "graphical_abstract_CES.docx"),
    ]
    for path in outputs:
        print(path)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
